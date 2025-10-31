import yaml
from docx import Document

class CloudFormationLoader(yaml.SafeLoader):
    """
    Custom YAML loader for handling CloudFormation intrinsic functions.
    """
    pass

# Add handlers for common CloudFormation intrinsic functions
for tag in ['!Ref', '!Sub', '!GetAtt', '!FindInMap', '!Join', '!Base64', '!Select', '!Split', '!If', '!And', '!Not', '!Equals', '!Or']:
    CloudFormationLoader.add_constructor(tag, lambda loader, node: node.value)

def parse_cloudformation_yaml(file_path):
    """
    Parse a CloudFormation YAML file and extract resources with their descriptions.
    """
    with open(file_path, 'r') as file:
        data = yaml.load(file, Loader=CloudFormationLoader)

    resources = data.get('Resources', {})
    resource_list = []

    for resource_name, resource_details in resources.items():
        resource_type = resource_details.get('Type', 'Unknown Type')
        description = resource_details.get('Properties', {}).get('Description', 'No Description Available')
        resource_list.append((resource_name, resource_type, description))
    
    return resource_list

def generate_doc(resources, output_path):
    """
    Generate a document listing all resources and their descriptions.
    """
    doc = Document()
    doc.add_heading('CloudFormation Resources', level=1)

    for resource_name, resource_type, description in resources:
        doc.add_heading(resource_name, level=2)
        doc.add_paragraph(f"Type: {resource_type}")
        doc.add_paragraph(f"Description: {description}")

    doc.save(output_path)

def main():
    input_file = 'template.yaml'  # Path to your YAML file
    output_file = 'CloudFormation_Resources.docx'  # Path to output document

    resources = parse_cloudformation_yaml(input_file)
    generate_doc(resources, output_file)
    print(f"Document created successfully: {output_file}")

if __name__ == '__main__':
    main()
